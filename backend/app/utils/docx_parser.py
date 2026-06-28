import io
import re  
from docx import Document

def extract_text_and_tables_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    content = []

   
    content.append("=== Text Content ===")
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if text and len(re.sub(r'[^a-zA-Z0-9\u0621-\u064A]', '', text)) > 0:
            content.append(text)

   
    if doc.tables:
        content.append("\n=== Tables ===")
        for i, table in enumerate(doc.tables):
            
            if not table.rows: continue
            
            content.append(f"\n[Table {i+1}]")
            for row_idx, row in enumerate(table.rows):
                
                row_data = [cell.text.strip().replace("\n", " ").replace("|", "-") for cell in row.cells]
                markdown_row = "| " + " | ".join(row_data) + " |"
                content.append(markdown_row)
                
                if row_idx == 0:
                    separator = "|" + "|".join(["---"] * len(row.cells)) + "|"
                    content.append(separator)

    return "\n".join(content)